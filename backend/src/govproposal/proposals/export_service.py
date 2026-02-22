"""Proposal document export service."""

import re
from datetime import datetime, timezone
from io import BytesIO
from typing import Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


class ProposalExportService:
    """Service for exporting proposals as formatted Word documents."""

    def generate_docx(self, proposal_data: dict, org_data: dict) -> BytesIO:
        """Generate a formatted Word document from proposal data.

        Args:
            proposal_data: Dict with title, sections, metadata
            org_data: Dict with org name, credentials

        Returns:
            BytesIO buffer containing the .docx file
        """
        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Configure heading styles
        for level in range(1, 4):
            heading_style = doc.styles[f"Heading {level}"]
            heading_style.font.name = "Calibri"
            heading_style.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
            if level == 1:
                heading_style.font.size = Pt(16)
                heading_style.font.bold = True
            elif level == 2:
                heading_style.font.size = Pt(14)
                heading_style.font.bold = True
            else:
                heading_style.font.size = Pt(12)
                heading_style.font.bold = True

        # --- Cover Page ---
        self._add_cover_page(doc, proposal_data, org_data)

        # --- Content Sections ---
        sections = [
            ("1.0", "Executive Summary", proposal_data.get("executive_summary")),
            ("2.0", "Technical Approach", proposal_data.get("technical_approach")),
            ("3.0", "Management Approach", proposal_data.get("management_approach")),
            ("4.0", "Past Performance", proposal_data.get("past_performance")),
            ("5.0", "Pricing Summary", proposal_data.get("pricing_summary")),
        ]

        for number, title, content in sections:
            doc.add_page_break()
            doc.add_heading(f"{number}  {title}", level=1)

            if content:
                self._add_markdown_content(doc, content)
            else:
                p = doc.add_paragraph("[This section has not been completed.]")
                p.runs[0].font.italic = True
                p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # --- Footer with page numbers ---
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        org_name = org_data.get("name", "")
        sol_num = proposal_data.get("solicitation_number", "")
        run.text = f"{org_name}  |  {sol_num}  |  PROPRIETARY"

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _add_cover_page(self, doc: Document, proposal_data: dict, org_data: dict) -> None:
        """Add a professional cover page."""
        # Spacer
        for _ in range(4):
            doc.add_paragraph("")

        # Organization name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(org_data.get("name", "Organization"))
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

        doc.add_paragraph("")

        # Proposal title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(proposal_data.get("title", "Proposal"))
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        doc.add_paragraph("")

        # Divider line
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("━" * 40)
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        run.font.size = Pt(12)

        doc.add_paragraph("")

        # Metadata table
        details = []
        if proposal_data.get("agency"):
            details.append(("Agency:", proposal_data["agency"]))
        if proposal_data.get("solicitation_number"):
            details.append(("Solicitation Number:", proposal_data["solicitation_number"]))
        if proposal_data.get("naics_code"):
            details.append(("NAICS Code:", proposal_data["naics_code"]))
        if proposal_data.get("due_date"):
            details.append(("Response Deadline:", proposal_data["due_date"]))
        if proposal_data.get("estimated_value"):
            details.append(("Estimated Value:", f"${proposal_data['estimated_value']:,.0f}"))

        if org_data.get("uei_number"):
            details.append(("UEI:", org_data["uei_number"]))
        if org_data.get("cage_code"):
            details.append(("CAGE Code:", org_data["cage_code"]))

        details.append(("Date Prepared:", datetime.now(timezone.utc).strftime("%B %d, %Y")))

        for label, value in details:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_label = p.add_run(f"{label} ")
            run_label.font.bold = True
            run_label.font.size = Pt(11)
            run_label.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run_value = p.add_run(str(value))
            run_value.font.size = Pt(11)
            run_value.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Footer notice
        for _ in range(3):
            doc.add_paragraph("")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("PROPRIETARY AND CONFIDENTIAL")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    def _add_markdown_content(self, doc: Document, text: str) -> None:
        """Parse markdown text and add it to the document with formatting."""
        lines = text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Skip empty lines
            if not stripped:
                i += 1
                continue

            # Headings
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            # Numbered lists
            elif re.match(r"^\d+\.\s", stripped):
                content = re.sub(r"^\d+\.\s", "", stripped)
                p = doc.add_paragraph(style="List Number")
                self._add_formatted_runs(p, content)
            # Bullet lists
            elif stripped.startswith("- ") or stripped.startswith("* "):
                content = stripped[2:]
                p = doc.add_paragraph(style="List Bullet")
                self._add_formatted_runs(p, content)
            # Regular paragraph
            else:
                # Collect consecutive non-special lines into one paragraph
                para_lines = [stripped]
                while i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if (
                        not next_line
                        or next_line.startswith("#")
                        or next_line.startswith("- ")
                        or next_line.startswith("* ")
                        or re.match(r"^\d+\.\s", next_line)
                    ):
                        break
                    para_lines.append(next_line)
                    i += 1

                full_text = " ".join(para_lines)
                p = doc.add_paragraph()
                self._add_formatted_runs(p, full_text)

            i += 1

    def _add_formatted_runs(self, paragraph, text: str) -> None:
        """Add text with bold and italic formatting to a paragraph."""
        # Pattern: **bold**, *italic*, ***bold italic***
        pattern = r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)"
        parts = re.split(pattern, text)

        i = 0
        while i < len(parts):
            part = parts[i]
            if part is None:
                i += 1
                continue

            # Check if this is a formatting match
            if part.startswith("***") and part.endswith("***"):
                run = paragraph.add_run(parts[i + 1])
                run.font.bold = True
                run.font.italic = True
                i += 2
            elif part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(parts[i + 2])
                run.font.bold = True
                i += 3
            elif part.startswith("*") and part.endswith("*") and len(part) > 1:
                run = paragraph.add_run(parts[i + 3])
                run.font.italic = True
                i += 4
            else:
                if part:
                    paragraph.add_run(part)
                i += 1
