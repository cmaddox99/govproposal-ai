"""Extract structured past-performance fields from an uploaded document.

Pipeline:
  1. Read the binary file from a path (already stored on disk).
  2. Convert to plain text based on content type (pdf, docx, txt, csv).
  3. Ask Claude to map the text to PastPerformanceCreate-shaped JSON.

The extraction never persists anything by itself — the caller decides whether
to create a record. The source file remains in storage either way; an orphaned
file is acceptable for now.
"""

from __future__ import annotations

import io
import json
import logging
from typing import Any

import anthropic

from govproposal.config import settings

logger = logging.getLogger(__name__)


MAX_EXTRACT_CHARS = 30_000


def extract_text(path: str, content_type: str | None, filename: str) -> str:
    """Extract plain text from a stored file. Returns empty string on failure."""
    lower = (filename or "").lower()

    if (content_type == "application/pdf") or lower.endswith(".pdf"):
        return _read_pdf(path)

    if (
        content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or lower.endswith(".docx")
    ):
        return _read_docx(path)

    if (content_type or "").startswith("text/") or lower.endswith((".txt", ".csv", ".md")):
        return _read_text(path)

    # Best-effort fallback: try text first, then PDF, then DOCX
    for reader in (_read_text, _read_pdf, _read_docx):
        try:
            text = reader(path)
            if text.strip():
                return text
        except Exception:
            continue
    return ""


def _read_text(path: str) -> str:
    with open(path, "rb") as fh:
        data = fh.read()
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _read_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.warning("pypdf not installed; cannot read PDF")
        return ""

    try:
        reader = PdfReader(path)
        chunks: list[str] = []
        for page in reader.pages:
            try:
                chunks.append(page.extract_text() or "")
            except Exception as exc:
                logger.warning("PDF page extraction failed: %s", exc)
        return "\n\n".join(chunks)
    except Exception as exc:
        logger.warning("Failed to read PDF %s: %s", path, exc)
        return ""


def _read_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed; cannot read DOCX")
        return ""

    try:
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as exc:
        logger.warning("Failed to read DOCX %s: %s", path, exc)
        return ""


EXTRACTION_SYSTEM_PROMPT = """You extract structured fields from past-performance documents for a government contracting platform.

You receive the plain text of a contract, CPARS report, project summary, or similar document. A single document may describe ONE past performance or SEVERAL distinct past performances (e.g. a references package listing multiple contracts). You produce a JSON object matching this schema:

{
  "records": [                         // one entry PER distinct past performance in the document
    {
      "contract_name": string,           // required — the project / contract title
      "agency": string | null,           // the contracting agency or customer name
      "contract_number": string | null,  // contract / task order number
      "contract_value": number | null,   // total contract value in USD (number only, no currency symbol)
      "period_of_performance_start": string | null,  // ISO 8601 date "YYYY-MM-DD"
      "period_of_performance_end": string | null,    // ISO 8601 date "YYYY-MM-DD"
      "description": string | null,      // 2-4 sentence summary of scope of work
      "contact_name": string | null,     // government technical point of contact
      "contact_email": string | null,
      "contact_phone": string | null,
      "performance_rating": string | null,  // CPARS-style rating: "Exceptional", "Very Good", "Satisfactory", "Marginal", "Unsatisfactory" — or null if unknown
      "confidence": {                    // 0.0 - 1.0 per field, your confidence in the value
        "contract_name": number,
        "agency": number,
        "contract_number": number,
        "contract_value": number,
        "period_of_performance_start": number,
        "period_of_performance_end": number,
        "description": number,
        "contact_name": number,
        "contact_email": number,
        "contact_phone": number,
        "performance_rating": number
      }
    }
  ]
}

Rules:
- Output ONLY the JSON object. No prose, no markdown fences, no commentary.
- Emit one records[] entry for EACH distinct contract / past performance described. Do not merge separate contracts into one entry, and do not split one contract into several.
- If a field cannot be found, set it to null and confidence to 0.
- For dates, prefer the explicit period of performance over award dates.
- For contract_value, sum line items if the document lists them; if a range is given, use the maximum.
- contract_name is required — if no clear title exists, build one from agency + scope (e.g. "DHA Medical Equipment Maintenance").
"""


async def extract_past_performance(text: str) -> list[dict[str, Any]]:
    """Run Claude over document text and return one dict per past performance found."""
    if not text or not text.strip():
        return []
    if not settings.anthropic_api_key:
        logger.info("Anthropic key not configured; skipping past-performance extraction")
        return []

    snippet = text[:MAX_EXTRACT_CHARS]
    client = anthropic.AsyncAnthropic(api_key=settings.anthropic_api_key)

    try:
        message = await client.messages.create(
            model=settings.anthropic_model,
            max_tokens=8192,
            system=EXTRACTION_SYSTEM_PROMPT,
            messages=[
                {
                    "role": "user",
                    "content": (
                        "Extract past-performance fields from this document. "
                        "Return only the JSON object as specified, with one "
                        "records[] entry per distinct past performance.\n\n"
                        f"<document>\n{snippet}\n</document>"
                    ),
                }
            ],
        )
    except anthropic.AuthenticationError:
        logger.error("Invalid Anthropic API key for past-performance extraction")
        return []
    except anthropic.RateLimitError:
        logger.warning("Anthropic rate limit reached during extraction")
        return []
    except Exception as exc:
        logger.error("Claude extraction error: %s", exc)
        return []

    try:
        raw = message.content[0].text
        if "```json" in raw:
            raw = raw.split("```json", 1)[1].split("```", 1)[0]
        elif "```" in raw:
            raw = raw.split("```", 1)[1].split("```", 1)[0]
        parsed = json.loads(raw.strip())
    except (json.JSONDecodeError, IndexError, AttributeError) as exc:
        logger.warning("Could not parse extraction JSON: %s", exc)
        return []

    # Expected shape: {"records": [...]}. Tolerate a bare list or a single
    # object (older prompt shape) so a drifting model response still works.
    if isinstance(parsed, dict) and isinstance(parsed.get("records"), list):
        records = parsed["records"]
    elif isinstance(parsed, list):
        records = parsed
    elif isinstance(parsed, dict):
        records = [parsed]
    else:
        return []

    return [r for r in records if isinstance(r, dict) and r.get("contract_name")]
